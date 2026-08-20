from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..styles import apply_font
from ..figures import add_logo


def add_certificate_page(document, data):

    document.add_page_break()

    for _ in range(3):
        document.add_paragraph()

    def center(text, size=12, bold=False, italic=False, underline=False):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        apply_font(r, size, bold, italic, underline)
        return p

    center("JAI HIND COLLEGE", 14, True)
    center("(Autonomous)", 11, True, True)
    center("MUMBAI, 400020 MAHARASHTRA", 12, True)

    document.add_paragraph()
    center("DEPARTMENT OF INFORMATION TECHNOLOGY", 14, True)

    document.add_paragraph()
    add_logo(document, "templates/jai_hind/logo.png", 3.2)

    document.add_paragraph()
    center("CERTIFICATE", 14, True, False, True)

    # Certificate content
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    r = p.add_run(
        f'This is to certify that the project entitled, '
        f'“{data["project_title"]}” is the bonafide work of '
        f'{data["student_name"]} bearing UID ({data["uid"]}) '
        f'submitted in partial fulfillment of the requirements for '
        f'the award of degree of {data["degree"]} from Jai Hind College '
        f'Autonomous (University of Mumbai).'
    )
    apply_font(r, 12)

    # Signature section
    document.add_paragraph("\n\n\n")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r = p.add_run("Internal Guide")
    apply_font(r, 12, True)

    r = p.add_run("\t\t\t\tCoordinator")
    apply_font(r, 12, True)

    document.add_paragraph("\n\n")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r = p.add_run("Date")
    apply_font(r, 12, True)

    r = p.add_run("\t\t\t\tExternal Examiner")
    apply_font(r, 12, True)

    document.add_paragraph("\n\n\n\n")

    center("College Seal", 12, True)