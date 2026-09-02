from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm
from ..styles import apply_font


def add_index_page(document, data):

    document.add_page_break()

    # Top spacing
    for _ in range(3):
        document.add_paragraph()

    # INDEX heading
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run("INDEX")
    apply_font(r, 16, bold=True)

    document.add_paragraph()

    rows = data.get("index", [])

    # ---------------------------------------------------------
    # 6 COLUMNS
    # Sr No | Level 1 | Level 2 | Particulars | Page No | Date
    # ---------------------------------------------------------

    table = document.add_table(rows=1, cols=6)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [
        Cm(1.0),   # Sr No
        Cm(1.0),   # Level 1
        Cm(1.3),   # Level 2
        Cm(8.2),   # Particulars
        Cm(1.8),   # Page No
        Cm(2.5),   # Date
    ]

    headers = [
        "Sr\nNo",
        "",
        "",
        "PARTICULARS",
        "PAGE\nNO.",
        "DATE"
    ]

    # ---------------------------------------------------------
    # HEADER
    # ---------------------------------------------------------

    for i, header in enumerate(headers):

        cell = table.rows[0].cells[i]

        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r = p.add_run(header)
        apply_font(r, 11, bold=True)

    # ---------------------------------------------------------
    # CONTENT ROWS
    # ---------------------------------------------------------

    for item in rows:

        cells = table.add_row().cells

        values = [
            item.get("sr", ""),
            item.get("level1", ""),
            item.get("level2", ""),
            item.get("title", ""),
            item.get("page", ""),
            item.get("date", ""),
        ]

        for i, value in enumerate(values):

            cells[i].width = widths[i]

            cells[i].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            r = p.add_run(str(value))

            apply_font(
                r,
                11,
                bold=item.get("bold", False)
            )

def add_figures_page(document, data):

    document.add_page_break()

    for _ in range(3):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LIST OF FIGURES")
    apply_font(r, 16, bold=True)

    document.add_paragraph()

    rows = data.get("figures_index", [])

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(1.5), Cm(11.3), Cm(2.5)]
    headers = ["Sr\nNo", "FIGURE TITLE", "PAGE\nNO."]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        apply_font(r, 11, bold=True)

    for item in rows:
        cells = table.add_row().cells
        values = [item.get("sr", ""), item.get("title", ""), item.get("page", "")]
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            apply_font(r, 11)