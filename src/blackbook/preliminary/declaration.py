from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..styles import apply_font


def add_declaration_page(document, data):

    # FORCE declaration onto a completely new page
    document.add_page_break()

    # Heading spacing
    for _ in range(4):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run("DECLARATION")
    apply_font(r, 14, True, False, True)

    document.add_paragraph("\n")

    paragraphs = [
        (
            f'I hereby declare that the project entitled, '
            f'“{data["project_title"]}” done at Jai Hind College, '
            f'has not been in any case duplicated to submit to any '
            f'other university for the award of any degree. To the '
            f'best of my knowledge other than me, no one has submitted '
            f'to any other university.'
        ),
        (
            f'The project is done in partial fulfillment of the '
            f'requirements for the award of degree of {data["degree"]} '
            f'to be submitted as final semester project as part of '
            f'our curriculum.'
        )
    ]

    for text in paragraphs:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = 18

        r = p.add_run(text)
        apply_font(r, 12)

    # Signature at lower-right area
    for _ in range(7):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r = p.add_run("Name and Signature of the Student")
    apply_font(r, 12, True)