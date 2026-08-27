from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def apply_font(run, size=12, bold=False, italic=False, underline=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    apply_font(run, size=12)
    return paragraph


def add_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    apply_font(run, size=16, bold=True)
    return paragraph


def add_section_heading(document, text):
    """14pt bold — numbered sections like 1.1, 3.4. Triggers new page in renderer."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    apply_font(run, size=14, bold=True)
    return paragraph


def add_subsection_heading(document, text):
    """12pt bold — numbered subsections like 1.3.1, 3.6.2. No page break."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    apply_font(run, size=12, bold=True)
    return paragraph


def add_label_heading(document, text):
    """14pt bold — unnumbered display labels like '1] SOFTWARE:'. 
    No page break, no index entry, purely visual weight."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    apply_font(run, size=14, bold=True)
    paragraph.paragraph_format.space_before = Pt(6)
    return paragraph


def add_bullet_square(document, text, bold=False, indent_cm=0.5):
    """☐ style bullet — category labels like 'Frontend Technologies:'"""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    run = paragraph.add_run("☐  ")
    apply_font(run, size=12, bold=True)
    run2 = paragraph.add_run(text)
    apply_font(run2, size=12, bold=bold)
    return paragraph


def add_bullet_arrow(document, text, indent_cm=1.2):
    """➢ style bullet — items nested under a ☐ category"""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    run = paragraph.add_run("➢  ")
    apply_font(run, size=12)
    run2 = paragraph.add_run(text)
    apply_font(run2, size=12)
    return paragraph


def add_bullet_round(document, text, indent_cm=0.75):
    """• style bullet — plain flat lists (e.g. Hardware section)"""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    run = paragraph.add_run("•  ")
    apply_font(run, size=12)
    run2 = paragraph.add_run(text)
    apply_font(run2, size=12)
    return paragraph


def add_bullet_with_bold_lead(document, bold_part, rest_text, indent_cm=0.75):
    """➢ bullet with bold lead-in term + regular continuation, e.g.
    '➢ Secure Authentication: Role-based login...'"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(indent_cm)

    bullet_run = paragraph.add_run("➢  ")
    apply_font(bullet_run, size=12)

    bold_run = paragraph.add_run(bold_part + ": ")
    apply_font(bold_run, size=12, bold=True)

    rest_run = paragraph.add_run(rest_text)
    apply_font(rest_run, size=12)

    return paragraph


def add_centered_text(document, text, size=12, bold=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    apply_font(run, size=size, bold=bold)
    return paragraph


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15