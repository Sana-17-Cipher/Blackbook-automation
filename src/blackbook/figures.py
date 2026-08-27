from pathlib import Path
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

def add_logo(document, logo_path, width_cm=3.2):
    """Add the institutional logo centered on the page."""

    logo_path = Path(logo_path)

    if not logo_path.exists():
        raise FileNotFoundError(f"Logo not found: {logo_path}")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.add_picture(
        str(logo_path),
        width=Cm(width_cm)
    )

    return paragraph

FIGURES_DIR = Path("src/blackbook/assets")

def insert_figure(document, fig_number, width_cm=14):
    chapter_num = fig_number.split(".")[0]
    chapter_folder = FIGURES_DIR / f"chapter{chapter_num}"

    matches = list(chapter_folder.glob(f"fig-{fig_number}*"))
    if not matches:
        raise FileNotFoundError(f"No image found for figure {fig_number} in {chapter_folder}")

    image_path = matches[0]

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_text = _get_caption_text(image_path)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(f"Figure {fig_number}: {caption_text}")
    caption_run.bold = True
    caption_run.font.size = Pt(12)
    caption_run.font.name = "Times New Roman"
    caption.paragraph_format.space_after = Pt(10)

    return paragraph


def _get_caption_text(image_path):
    name = image_path.stem
    parts = name.split("_", 1)
    if len(parts) > 1:
        return parts[1].replace("_", " ").replace("-", " ").title()
    return ""