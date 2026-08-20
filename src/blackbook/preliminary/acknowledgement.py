from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..styles import apply_font


def add_acknowledgement_page(document, data):

    document.add_page_break()

    for _ in range(3):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ACKNOWLEDGEMENT")
    apply_font(r, 14, True, False, True)

    document.add_paragraph()

    text = data.get("acknowledgement", "")

    for paragraph_text in text.split("\n\n"):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(paragraph_text.strip())
        apply_font(r, 12)