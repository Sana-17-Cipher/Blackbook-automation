from pathlib import Path
import re

from docx.enum.text import WD_BREAK
from docx.shared import Pt

from .structure_detector import (
    detect_structure,
    CHAPTER,
    SECTION,
    SUBSECTION,
    SUBHEADING,
    BULLET,
    NUMBERED,
    BODY,
)


def add_page_break(document):
    """Add a page break."""
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def clean_inline_markdown(text):
    """
    Remove Markdown formatting markers, in case any slip through
    (e.g. the user pasted something with **bold** in it by habit).
    """
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text


def add_heading(document, text, level):
    """
    Blackbook heading hierarchy:

    Level 1 = 16 pt bold
    Level 2 = 14 pt bold
    Level 3 = 12 pt bold
    """
    text = clean_inline_markdown(text)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True

    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)

    run.font.name = "Times New Roman"

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)

    return paragraph


def add_body_paragraph(document, text):
    """Add normal 12 pt body text."""
    text = clean_inline_markdown(text)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    paragraph.paragraph_format.first_line_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15

    return paragraph


def add_bullet(document, text, bold_prefix=None):
    """
    Add a 12 pt bullet.

    If bold_prefix is given (e.g. "React.js" for the bullet
    "React.js: A powerful JavaScript library..."), that leading term
    is rendered bold and the rest of the text stays normal weight,
    matching the reference formatting.
    """
    paragraph = document.add_paragraph(style="List Bullet")

    if bold_prefix:
        remainder = text
        prefix_marker = f"{bold_prefix}:"
        if remainder.startswith(prefix_marker):
            remainder = remainder[len(prefix_marker):]

        bold_run = paragraph.add_run(clean_inline_markdown(bold_prefix) + ":")
        bold_run.bold = True
        bold_run.font.size = Pt(12)
        bold_run.font.name = "Times New Roman"

        rest_run = paragraph.add_run(clean_inline_markdown(remainder))
        rest_run.font.size = Pt(12)
        rest_run.font.name = "Times New Roman"
    else:
        text = clean_inline_markdown(text)
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    paragraph.paragraph_format.space_after = Pt(4)

    return paragraph


def add_subheading(document, text):
    """
    Add a bold numbered sub-heading inside a chapter, e.g.
    "1. Frontend Technologies" -- bold, 12 pt, slightly indented,
    no page break of its own.
    """
    text = clean_inline_markdown(text)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Pt(18)

    return paragraph


def add_numbered_item(document, text):
    """Add a 12 pt numbered item."""
    text = clean_inline_markdown(text)

    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    paragraph.paragraph_format.space_after = Pt(4)

    return paragraph


# ---------------------------------------------------------------------
# NEW: render directly from plain text -- no Markdown step required.
# ---------------------------------------------------------------------

def render_plain_text(document, text):
    """
    Render a blackbook chapter directly from plain, unformatted text.

    No '#', '##', '###' needed anywhere. Structure is detected
    automatically from the numbering convention:

        1.1 Background      -> section    (new page, 14pt bold)
        1.3.1 Purpose        -> subsection (same page, 12pt bold)
        1. Automated profiling -> numbered list item
        INTRODUCTION          -> chapter   (16pt bold)

    Page-break rule (same as the Markdown renderer):
        - the very first heading in the file does NOT get an extra
          page break (the preliminary pages already ended on their
          own page)
        - every SECTION after the first one gets a page break before it
        - CHAPTER and SUBSECTION headings never force their own break
    """
    tokens = detect_structure(text)
    _render_tokens(document, tokens)


def render_plain_text_file(document, text_file):  # e.g. "introduction.txt"
    """Same as render_plain_text, but reads the content from a file."""
    text_path = Path(text_file)

    if not text_path.exists():
        raise FileNotFoundError(f"Text file not found: {text_path}")

    text = text_path.read_text(encoding="utf-8")
    render_plain_text(document, text)


def _render_tokens(document, tokens):
    first_heading_seen = False

    for token in tokens:
        if token.type == CHAPTER:
            # A chapter heading never triggers a page break itself;
            # whatever chapter/section came before it already
            # ended on its own page (via the section rule below).
            add_heading(document, token.text, 1)
            first_heading_seen = True

        elif token.type == SECTION:
            if first_heading_seen:
                add_page_break(document)
            first_heading_seen = True
            add_heading(document, token.text, 2)

        elif token.type == SUBSECTION:
            add_heading(document, token.text, 3)
            first_heading_seen = True

        elif token.type == SUBHEADING:
            add_subheading(document, token.text)
            first_heading_seen = True

        elif token.type == BULLET:
            add_bullet(document, token.text, bold_prefix=token.bold_prefix)

        elif token.type == NUMBERED:
            add_numbered_item(document, token.text)

        elif token.type == BODY:
            add_body_paragraph(document, token.text)


# ---------------------------------------------------------------------
# Original Markdown-based renderer, kept so any .md files you already
# have keep working without changes.
# ---------------------------------------------------------------------

def render_markdown(document, markdown_file):
    """Render a blackbook Markdown chapter (legacy path)."""
    markdown_path = Path(markdown_file)

    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {markdown_path}")

    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    current_paragraph = []
    first_section = True

    def flush_paragraph():
        if current_paragraph:
            paragraph_text = " ".join(
                line.strip() for line in current_paragraph
            ).strip()
            if paragraph_text:
                add_body_paragraph(document, paragraph_text)
            current_paragraph.clear()

    for raw_line in lines:
        line = raw_line.strip()

        if not line:
            flush_paragraph()
            continue

        if re.match(r"^#\s+", line):
            flush_paragraph()
            heading_text = re.sub(r"^#\s+", "", line).strip()
            if heading_text.upper() == "INTRODUCTION":
                continue
            add_heading(document, heading_text, 1)
            continue

        if re.match(r"^##\s+", line):
            flush_paragraph()
            if not first_section:
                add_page_break(document)
            first_section = False
            heading_text = re.sub(r"^##\s+", "", line).strip()
            add_heading(document, heading_text, 2)
            continue

        if re.match(r"^###\s+", line):
            flush_paragraph()
            heading_text = re.sub(r"^###\s+", "", line).strip()
            add_heading(document, heading_text, 3)
            continue

        if re.match(r"^[-*]\s+", line):
            flush_paragraph()
            bullet_text = re.sub(r"^[-*]\s+", "", line)
            add_bullet(document, bullet_text)
            continue

        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            numbered_text = re.sub(r"^\d+\.\s+", "", line)
            add_numbered_item(document, numbered_text)
            continue

        current_paragraph.append(clean_inline_markdown(line))

    flush_paragraph()