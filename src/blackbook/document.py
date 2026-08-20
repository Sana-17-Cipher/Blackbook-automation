from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .config import BlackbookConfig


def set_page_border(section, config):
    """Add the rectangular border around every page."""

    sect_pr = section._sectPr

    page_borders = OxmlElement("w:pgBorders")
    page_borders.set(qn("w:offsetFrom"), "page")

    for edge in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(config.border_size))
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), config.border_color)

        page_borders.append(border)

    sect_pr.append(page_borders)


def add_page_number(section):
    """Add an automatic page number to the footer."""

    footer = section.footer
    paragraph = footer.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    field_instruction = OxmlElement("w:instrText")
    field_instruction.set(qn("xml:space"), "preserve")
    field_instruction.text = "PAGE"

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(field_instruction)
    run._r.append(field_end)

    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def create_document(config=None):
    """Create the base blackbook Word document."""

    if config is None:
        config = BlackbookConfig()

    document = Document()

    section = document.sections[0]

    # A4 page
    section.page_width = Cm(config.page_width_cm)
    section.page_height = Cm(config.page_height_cm)

    # Margins
    section.top_margin = Cm(config.top_margin_cm)
    section.bottom_margin = Cm(config.bottom_margin_cm)
    section.left_margin = Cm(config.left_margin_cm)
    section.right_margin = Cm(config.right_margin_cm)

    # Page border
    if config.border_enabled:
        set_page_border(section, config)

    # Page number
    add_page_number(section)

    # Default body style
    normal_style = document.styles["Normal"]

    normal_style.font.name = config.font_name
    normal_style.font.size = Pt(config.body_font_size)

    normal_style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        config.font_name
    )

    normal_style.paragraph_format.line_spacing = (
        config.body_line_spacing
    )

    normal_style.paragraph_format.space_after = Pt(
        config.paragraph_spacing_after_pt
    )

    normal_style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
    )

    return document