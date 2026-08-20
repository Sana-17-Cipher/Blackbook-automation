from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..styles import apply_font


def add_abstract_page(document, data):

    document.add_page_break()

    # Heading spacing
    for _ in range(3):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run("ABSTRACT")
    apply_font(r, 16, bold=True)

    document.add_paragraph()

    # Remove unwanted line breaks / extra spaces
    text = " ".join(data.get("abstract", "").split())

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = 0

    r = p.add_run(text)
    apply_font(r, 12, bold=True)