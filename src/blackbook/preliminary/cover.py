from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from ..figures import add_logo
from ..styles import apply_font


def add_centered(document, text, size=12, bold=False, italic=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    apply_font(run, size=size, bold=bold, italic=italic)

    return paragraph


def add_cover_page(document, data):
    """Generate the Jai Hind College blackbook cover page."""

    # Top spacing
    for _ in range(3):
        document.add_paragraph()

    add_centered(
        document,
        data["project_title"],
        size=18,
        bold=True
    )

    add_centered(
        document,
        data["project_subtitle"],
        size=15,
        bold=True
    )

    add_centered(
        document,
        "A Project Report",
        size=14,
        bold=True
    )

    add_centered(
        document,
        "Submitted in partial fulfilment of the",
        size=12
    )

    add_centered(
        document,
        data["semester_requirement"],
        size=12
    )

    add_centered(
        document,
        data["degree"],
        size=14,
        bold=True
    )

    add_centered(
        document,
        "BY",
        size=13,
        bold=True
    )

    add_centered(
        document,
        f"Name: {data['student_name']}",
        size=13,
        bold=True
    )

    add_centered(
        document,
        f"UID: {data['uid']}",
        size=13,
        bold=True
    )

    add_centered(
        document,
        "Under the esteemed guidance of",
        size=13,
        bold=True
    )

    add_centered(
        document,
        data["guide"],
        size=13,
        bold=True
    )

    add_centered(
        document,
        "Co-ordinator",
        size=13,
        bold=True
    )

    add_centered(
        document,
        "and",
        size=13,
        bold=True
    )

    add_centered(
        document,
        data["coordinator"],
        size=13,
        bold=True
    )

    add_centered(
        document,
        "Assistant Professor",
        size=13,
        bold=True
    )

    add_centered(
        document,
        "DEPARTMENT OF INFORMATION TECHNOLOGY",
        size=15,
        bold=True
    )

    add_centered(
        document,
        "JAI HIND COLLEGE",
        size=15,
        bold=True
    )
    add_logo(
    document,
    "templates/jai_hind/logo.png",
    width_cm=3.2
)
    # Logo will be added here in the next step.
    document.add_paragraph()

    add_centered(
        document,
        "(Autonomous)",
        size=12,
        bold=True,
        italic=True
    )

    add_centered(
        document,
        "MUMBAI, 400020",
        size=12,
        bold=True
    )

    add_centered(
        document,
        "MAHARASHTRA",
        size=12,
        bold=True
    )

    add_centered(
        document,
        data["academic_year"],
        size=12
    )